import docx

# create doc
doc = docx.Document("demo.docx")
print(doc.paragraphs)
print(doc.paragraphs[0].text) # see string in 1st paragraph

# paragraph font styles (bold, italic, underline)
p = doc.paragraphs[1]
print(p.runs)

print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)

# check style bold, italic, underline
print("check bold: " + str(p.runs[1].bold))

# set style
p.runs[3].underline = True

# edit text
p.runs[3].text = "italic and underline"

doc.save("demo2.docx")

# style (heading, normal)
print(p.style)
p.style = "Title"


# create empty doc
d = docx.Document()
d.add_paragraph("hello this is just some text")
d.add_paragraph("this is more text")
doc.save("demo3.docx")

p = d.paragraphs[0]
p.add_run("this is a new run")
print(p.runs)

# gettext from document
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return "\n".join(fullText)

print(getText("demo.docx"))
